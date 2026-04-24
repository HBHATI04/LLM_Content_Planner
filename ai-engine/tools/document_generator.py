"""
document_generator.py
Generates PDF and DOCX files from markdown-like AI response text.
Returns the file path on success, None on failure.
"""

import os
import re
import time
from pathlib import Path


def _clean_text(text: str) -> str:
    """Strip markdown bold/italic markers for plain-text renderers."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)   # **bold**
    text = re.sub(r"\*(.*?)\*", r"\1", text)         # *italic*
    text = re.sub(r"__(.*?)__", r"\1", text)          # __bold__
    return text.strip()


def _parse_blocks(text: str) -> list[dict]:
    """
    Parse text into blocks: heading, bullet, or paragraph.
    Returns list of {"type": ..., "text": ...}
    """
    blocks = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            blocks.append({"type": "h3", "text": stripped[4:]})
        elif stripped.startswith("## "):
            blocks.append({"type": "h2", "text": stripped[3:]})
        elif stripped.startswith("# "):
            blocks.append({"type": "h1", "text": stripped[2:]})
        elif stripped.startswith(("- ", "* ", "• ")):
            blocks.append({"type": "bullet", "text": stripped[2:]})
        elif re.match(r"^\d+\.\s", stripped):
            blocks.append({"type": "numbered", "text": re.sub(r"^\d+\.\s", "", stripped)})
        else:
            blocks.append({"type": "paragraph", "text": stripped})
    return blocks


def generate_pdf(content: str, title: str = "AI Generated Report") -> str | None:
    """Generate a PDF and return the file path."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        )

        os.makedirs("generated_docs", exist_ok=True)
        timestamp = int(time.time())
        file_path = f"generated_docs/doc_{timestamp}.pdf"

        doc = SimpleDocTemplate(
            file_path,
            pagesize=letter,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=22,
            textColor=colors.HexColor("#1e3a5f"),
            spaceAfter=16,
        )
        h1_style = ParagraphStyle(
            "CustomH1",
            parent=styles["Heading1"],
            fontSize=16,
            textColor=colors.HexColor("#1e3a5f"),
            spaceBefore=14,
            spaceAfter=6,
        )
        h2_style = ParagraphStyle(
            "CustomH2",
            parent=styles["Heading2"],
            fontSize=13,
            textColor=colors.HexColor("#2c5282"),
            spaceBefore=10,
            spaceAfter=4,
        )
        h3_style = ParagraphStyle(
            "CustomH3",
            parent=styles["Heading3"],
            fontSize=11,
            textColor=colors.HexColor("#2d3748"),
            spaceBefore=8,
            spaceAfter=3,
        )
        body_style = ParagraphStyle(
            "CustomBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=16,
            spaceAfter=6,
        )
        bullet_style = ParagraphStyle(
            "CustomBullet",
            parent=styles["Normal"],
            fontSize=10,
            leading=16,
            leftIndent=20,
            spaceAfter=3,
        )

        story = []
        story.append(Paragraph(_clean_text(title), title_style))
        story.append(Spacer(1, 12))

        blocks = _parse_blocks(content)
        bullet_group = []

        def flush_bullets():
            if bullet_group:
                items = [ListItem(Paragraph(_clean_text(b), bullet_style), leftIndent=20)
                         for b in bullet_group]
                story.append(ListFlowable(items, bulletType="bullet", start="•"))
                story.append(Spacer(1, 4))
                bullet_group.clear()

        for block in blocks:
            if block["type"] != "bullet":
                flush_bullets()

            if block["type"] == "h1":
                story.append(Paragraph(_clean_text(block["text"]), h1_style))
            elif block["type"] == "h2":
                story.append(Paragraph(_clean_text(block["text"]), h2_style))
            elif block["type"] == "h3":
                story.append(Paragraph(_clean_text(block["text"]), h3_style))
            elif block["type"] == "bullet":
                bullet_group.append(block["text"])
            elif block["type"] == "numbered":
                story.append(Paragraph(f"• {_clean_text(block['text'])}", bullet_style))
            elif block["type"] == "paragraph":
                story.append(Paragraph(_clean_text(block["text"]), body_style))

        flush_bullets()
        doc.build(story)

        print(f"[DocGen] PDF saved: {file_path}")
        return file_path

    except Exception as e:
        print(f"[DocGen PDF ERROR] {e}")
        import traceback; traceback.print_exc()
        return None


def generate_docx(content: str, title: str = "AI Generated Report") -> str | None:
    """Generate a DOCX and return the file path."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        os.makedirs("generated_docs", exist_ok=True)
        timestamp = int(time.time())
        file_path = f"generated_docs/doc_{timestamp}.docx"

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # Title
        title_para = doc.add_heading(_clean_text(title), level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title_para.runs[0]
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

        doc.add_paragraph()  # spacer

        blocks = _parse_blocks(content)
        for block in blocks:
            btype = block["type"]
            text = _clean_text(block["text"])

            if btype == "h1":
                h = doc.add_heading(text, level=1)
                h.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
            elif btype == "h2":
                h = doc.add_heading(text, level=2)
                h.runs[0].font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
            elif btype == "h3":
                doc.add_heading(text, level=3)
            elif btype in ("bullet", "numbered"):
                doc.add_paragraph(text, style="List Bullet")
            elif btype == "paragraph":
                p = doc.add_paragraph(text)
                p.runs[0].font.size = Pt(11) if p.runs else None

        doc.save(file_path)
        print(f"[DocGen] DOCX saved: {file_path}")
        return file_path

    except Exception as e:
        print(f"[DocGen DOCX ERROR] {e}")
        import traceback; traceback.print_exc()
        return None